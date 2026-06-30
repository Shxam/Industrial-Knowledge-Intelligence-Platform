"""
Document loader for multiple file formats
Supports: PDF, XLSX, DOCX, Images
"""
import fitz  # PyMuPDF
import openpyxl
from docx import Document as DocxDocument
from minio import Minio
from minio.error import S3Error
from pathlib import Path
from typing import Dict, Any, Optional, BinaryIO
import uuid
from datetime import datetime

from app.core.config import settings
from app.core.logging import logger


class DocumentLoader:
    """Multi-format document loader with MinIO storage"""
    
    def __init__(self):
        """Initialize document loader with MinIO client"""
        self.minio_client = Minio(
            settings.MINIO_ENDPOINT,
            access_key=settings.MINIO_ACCESS_KEY,
            secret_key=settings.MINIO_SECRET_KEY,
            secure=settings.MINIO_SECURE
        )
        self._ensure_bucket()
    
    def _ensure_bucket(self):
        """Ensure MinIO bucket exists"""
        try:
            if not self.minio_client.bucket_exists(settings.MINIO_BUCKET):
                self.minio_client.make_bucket(settings.MINIO_BUCKET)
                logger.info(f"Created MinIO bucket: {settings.MINIO_BUCKET}")
        except S3Error as e:
            logger.error(f"Error checking/creating bucket: {e}")
    
    def load_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from PDF file
        
        Returns dict with text, metadata, and page info
        """
        logger.info(f"Loading PDF: {file_path}")
        
        try:
            doc = fitz.open(file_path)
            
            # Extract text by page
            pages = []
            full_text = ""
            
            for page_num, page in enumerate(doc, start=1):
                page_text = page.get_text()
                pages.append({
                    "page_number": page_num,
                    "text": page_text,
                    "char_count": len(page_text)
                })
                full_text += f"\n\n--- Page {page_num} ---\n\n{page_text}"
            
            metadata = {
                "page_count": len(doc),
                "has_images": any(page.get_images() for page in doc),
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "created": doc.metadata.get("creationDate", "")
            }
            
            doc.close()
            
            logger.info(f"Extracted {len(pages)} pages from PDF")
            
            return {
                "text": full_text,
                "pages": pages,
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error loading PDF: {e}")
            raise
    
    def load_xlsx(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from Excel file
        
        Converts tables to text format
        """
        logger.info(f"Loading XLSX: {file_path}")
        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            sheets = []
            full_text = ""
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Convert sheet to text
                sheet_text = f"Sheet: {sheet_name}\n\n"
                
                for row in sheet.iter_rows(values_only=True):
                    # Skip empty rows
                    if any(cell is not None for cell in row):
                        row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                        sheet_text += row_text + "\n"
                
                sheets.append({
                    "sheet_name": sheet_name,
                    "text": sheet_text,
                    "row_count": sheet.max_row,
                    "column_count": sheet.max_column
                })
                
                full_text += f"\n\n--- {sheet_name} ---\n\n{sheet_text}"
            
            workbook.close()
            
            metadata = {
                "sheet_count": len(sheets),
                "sheets": [s["sheet_name"] for s in sheets]
            }
            
            logger.info(f"Extracted {len(sheets)} sheets from XLSX")
            
            return {
                "text": full_text,
                "sheets": sheets,
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error loading XLSX: {e}")
            raise
    
    def load_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from Word document
        """
        logger.info(f"Loading DOCX: {file_path}")
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract paragraphs
            paragraphs = []
            full_text = ""
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
                    full_text += para.text + "\n\n"
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_text = ""
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    table_text += row_text + "\n"
                tables.append(table_text)
                full_text += f"\n{table_text}\n"
            
            metadata = {
                "paragraph_count": len(paragraphs),
                "table_count": len(tables)
            }
            
            logger.info(f"Extracted {len(paragraphs)} paragraphs and {len(tables)} tables from DOCX")
            
            return {
                "text": full_text,
                "paragraphs": paragraphs,
                "tables": tables,
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error loading DOCX: {e}")
            raise
    
    def load_file(self, file_path: str, file_type: Optional[str] = None) -> Dict[str, Any]:
        """
        Load file based on extension or provided type
        
        Returns unified dict with text and metadata
        """
        if file_type is None:
            file_type = Path(file_path).suffix.lower()
        
        loaders = {
            ".pdf": self.load_pdf,
            ".xlsx": self.load_xlsx,
            ".xls": self.load_xlsx,
            ".docx": self.load_docx,
            ".doc": self.load_docx
        }
        
        loader = loaders.get(file_type)
        if not loader:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        return loader(file_path)
    
    def save_to_minio(
        self,
        file_data: BinaryIO,
        filename: str,
        content_type: str = "application/octet-stream"
    ) -> str:
        """
        Save file to MinIO object storage
        
        Returns the object name (key)
        """
        # Generate unique object name
        object_name = f"{uuid.uuid4()}_{filename}"
        
        try:
            # Get file size
            file_data.seek(0, 2)  # Seek to end
            file_size = file_data.tell()
            file_data.seek(0)  # Seek back to start
            
            # Upload to MinIO
            self.minio_client.put_object(
                settings.MINIO_BUCKET,
                object_name,
                file_data,
                length=file_size,
                content_type=content_type
            )
            
            logger.info(f"Uploaded file to MinIO: {object_name}")
            
            return object_name
            
        except S3Error as e:
            logger.error(f"Error uploading to MinIO: {e}")
            raise
    
    def get_from_minio(self, object_name: str) -> bytes:
        """
        Retrieve file from MinIO
        
        Returns file content as bytes
        """
        try:
            response = self.minio_client.get_object(
                settings.MINIO_BUCKET,
                object_name
            )
            data = response.read()
            response.close()
            response.release_conn()
            
            logger.info(f"Retrieved file from MinIO: {object_name}")
            
            return data
            
        except S3Error as e:
            logger.error(f"Error retrieving from MinIO: {e}")
            raise
    
    def delete_from_minio(self, object_name: str):
        """Delete file from MinIO"""
        try:
            self.minio_client.remove_object(
                settings.MINIO_BUCKET,
                object_name
            )
            logger.info(f"Deleted file from MinIO: {object_name}")
            
        except S3Error as e:
            logger.error(f"Error deleting from MinIO: {e}")
            raise


# Global loader instance
document_loader = DocumentLoader()
